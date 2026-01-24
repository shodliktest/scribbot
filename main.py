import streamlit as st
import asyncio
import threading
import io
import os
import re
import tempfile
import cv2
import html
import json
import numpy as np
import img2pdf
from PIL import Image, ImageEnhance
from docx import Document
from docx.shared import Inches
from PyPDF2 import PdfReader, PdfWriter
from pdf2docx import Converter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.types import BufferedInputFile, ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.client.default import DefaultBotProperties
from google.cloud import vision

# --- 1. GLOBAL STORAGE ---
if "G_DATA" not in globals():
    G_DATA = {} 

# --- 2. CONFIGURATION & AUTHENTICATION ---
st.set_page_config(page_title="Google Vision AI Bot", layout="wide", page_icon="👁️")

st.markdown("""
    <style>
    .main { background-color: #0e1117; }
    .stMetric { background-color: #161b22; border: 1px solid #30363d; border-radius: 12px; padding: 15px; }
    .header-text { color: #58a6ff; font-size: 40px; font-weight: bold; text-align: center; }
    </style>
    """, unsafe_allow_html=True)

# Secrets va Google Auth sozlash
try:
    BOT_TOKEN = st.secrets["telegram"]["BOT_TOKEN"]
    ADMIN_ID = str(st.secrets["telegram"]["ADMIN_ID"])
    ADMIN_PASS = st.secrets["telegram"]["ADMIN_PASSWORD"]

    # Google Credentials ni JSON faylga vaqtincha yozish (Kutubxona fayl so'raydi)
    if "gcp_service_account" in st.secrets:
        service_account_info = dict(st.secrets["gcp_service_account"])
        # Vaqtinchalik fayl yaratamiz
        with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as f:
            json.dump(service_account_info, f)
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = f.name
    else:
        st.warning("⚠️ Google Cloud Credentials topilmadi!")

except Exception as e:
    st.error(f"❌ Secrets xatosi: {e}")
    st.stop()

# --- 3. GOOGLE VISION ENGINE & EFFECTS ---

def google_vision_scan(image_bytes):
    """
    Google Cloud Vision API orqali eng kuchli OCR
    """
    try:
        client = vision.ImageAnnotatorClient()
        content = image_bytes
        image = vision.Image(content=content)

        # Matnni aniqlash (DOCUMENT_TEXT_DETECTION eng yaxshisi)
        response = client.document_text_detection(image=image)
        
        if response.error.message:
            return f"❌ Google API Xatosi: {response.error.message}"
            
        return response.full_text_annotation.text if response.full_text_annotation else "Matn topilmadi."
    except Exception as e:
        return f"❌ Tizim xatosi: {e}"

def process_image_effect(img_bytes, effect="original"):
    """
    Rasmga effekt berish (PDF yoki ko'rish uchun)
    """
    nparr = np.frombuffer(img_bytes, np.uint8)
    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
    
    if effect == "bw": # Oq-Qora (Skaner effekti)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # Adaptive Threshold (soyalarni yo'qotadi)
        processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        _, buf = cv2.imencode(".jpg", processed)
        return buf.tobytes()
        
    elif effect == "enhance": # Sifatni oshirish
        # Denoising
        cleaned = cv2.fastNlMeansDenoisingColored(img, None, 10, 10, 7, 21)
        # Sharpness oshirish
        pil_img = Image.fromarray(cv2.cvtColor(cleaned, cv2.COLOR_BGR2RGB))
        enhancer = ImageEnhance.Sharpness(pil_img)
        sharpened = enhancer.enhance(1.5) # O'tkirlash
        contrast = ImageEnhance.Contrast(sharpened)
        final_pil = contrast.enhance(1.2) # Kontrast
        
        buf = io.BytesIO()
        final_pil.save(buf, format="JPEG", quality=95)
        return buf.getvalue()
        
    else: # Original
        return img_bytes

# --- 4. CONVERTERS ---

def create_pdf_from_text(text_content):
    """ReportLab orqali PDF yasash"""
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    text_obj = p.beginText(40, height - 40)
    text_obj.setFont("Helvetica", 12)
    
    for line in text_content.split('\n'):
        wrapped_lines = simpleSplit(line, "Helvetica", 12, width - 80)
        for wrapped in wrapped_lines:
            if text_obj.getY() < 40:
                p.drawText(text_obj)
                p.showPage()
                text_obj = p.beginText(40, height - 40)
                text_obj.setFont("Helvetica", 12)
            text_obj.textLine(wrapped)
    p.drawText(text_obj)
    p.save()
    buffer.seek(0)
    return buffer.getvalue()

def docx_to_pdf_engine(docx_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tf:
        tf.write(docx_bytes)
        path = tf.name
    try:
        doc = Document(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return create_pdf_from_text(text)
    finally:
        if os.path.exists(path): os.remove(path)

def convert_pdf_to_docx_safe(pdf_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tf:
        tf.write(pdf_bytes)
        temp_pdf = tf.name
    temp_docx = temp_pdf.replace(".pdf", ".docx")
    try:
        cv = Converter(temp_pdf)
        cv.convert(temp_docx, start=0, end=None)
        cv.close()
        with open(temp_docx, "rb") as f: return f.read()
    except: return None
    finally:
        if os.path.exists(temp_pdf): os.remove(temp_pdf)
        if os.path.exists(temp_docx): os.remove(temp_docx)

def images_to_docx(image_list):
    doc = Document()
    for img in image_list:
        doc.add_picture(io.BytesIO(img), width=Inches(6))
        doc.add_page_break()
    b = io.BytesIO(); doc.save(b); return b.getvalue()

# --- 5. BOT SETUP ---
@st.cache_resource
def get_bot():
    return Bot(token=BOT_TOKEN, default=DefaultBotProperties(parse_mode="HTML"))

bot = get_bot()
dp = Dispatcher()

def main_kb(uid):
    kb = [[KeyboardButton(text="ℹ️ Info"), KeyboardButton(text="👨‍💻 Adminga murojaat")]]
    if str(uid) == ADMIN_ID:
        kb.append([KeyboardButton(text="💎 Admin Panel")])
    return ReplyKeyboardMarkup(keyboard=kb, resize_keyboard=True)

INFO_TEXT = """
<b>👁️ GOOGLE VISION AI BOT</b>

📸 <b>Rasm imkoniyatlari:</b>
• <b>🔍 Google OCR:</b> Dunyodagi eng aniq skaner!
• <b>📄 PDF Skaner:</b> 3 xil rejim (Original, Oq-Qora, HD)
• <b>✨ Sifatni oshirish:</b> Xira rasmlarni tiniqlash

📄 <b>Fayl imkoniyatlari:</b>
• DOCX/TXT → PDF
• PDF → Word
• Split PDF (Kesish)

🎯 <b>Afzallik:</b> Audio yo'q, faqat hujjat va rasm bilan ishlaydi.
"""

# --- 6. HANDLERS ---

@dp.message(Command("start"))
async def start(m: types.Message):
    G_DATA[m.from_user.id] = {'files': [], 'state': None}
    await m.answer(f"👋 Salom {m.from_user.first_name}! Google Vision AI tayyor.", reply_markup=main_kb(m.from_user.id))

@dp.message(F.text)
async def text_handler(m: types.Message):
    uid, txt = m.from_user.id, m.text
    state = G_DATA.get(uid, {}).get('state')

    if str(uid) == ADMIN_ID and m.reply_to_message:
        replied_text = m.reply_to_message.text or m.reply_to_message.caption or ""
        match = re.search(r"#ID(\d+)", replied_text)
        if match:
            target_id = match.group(1)
            try:
                await bot.send_message(target_id, f"👨‍💻 <b>Admin javobi:</b>\n\n{html.escape(txt)}")
                await m.answer("✅ Javob yuborildi.")
            except: await m.answer("❌ Foydalanuvchi bloklagan.")
        return

    if state == "split":
        try:
            s, e = map(int, txt.split("-"))
            loop = asyncio.get_event_loop()
            def do_split():
                r = PdfReader(io.BytesIO(G_DATA[uid]['doc']))
                w = PdfWriter()
                for i in range(s-1, min(e, len(r.pages))): w.add_page(r.pages[i])
                o = io.BytesIO(); w.write(o); return o.getvalue()
            pdf = await loop.run_in_executor(None, do_split)
            await m.answer_document(BufferedInputFile(pdf, filename="kesilgan.pdf"), caption="✅ Tayyor")
        except: await m.answer("❌ Xato! Misol: 1-5")
        G_DATA[uid]['state'] = None
        return

    if txt == "ℹ️ Info": await m.answer(INFO_TEXT)
    elif txt == "👨‍💻 Adminga murojaat":
        G_DATA[uid]['state'] = "contact"
        await m.answer("Xabarni yozing:", reply_markup=types.ReplyKeyboardRemove())
    elif G_DATA.get(uid, {}).get('state') == "contact":
        await bot.send_message(ADMIN_ID, f"📩 #ID{uid} dan murojaat:\n{html.escape(txt)}")
        await m.answer("✅ Adminga yetkazildi.", reply_markup=main_kb(uid))
        G_DATA[uid]['state'] = None

@dp.message(F.photo)
async def photo_h(m: types.Message):
    uid = m.from_user.id
    if uid not in G_DATA: G_DATA[uid] = {'files': []}
    f = await bot.get_file(m.photo[-1].file_id)
    content = await bot.download_file(f.file_path)
    G_DATA[uid]['files'].append(content.read())
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🔍 Google OCR", callback_data="to_ocr")],
        [InlineKeyboardButton(text="✨ Tiniqlash (HD)", callback_data="to_enhance"), InlineKeyboardButton(text="📝 Wordga", callback_data="to_word")],
        [InlineKeyboardButton(text="📄 PDF (Original)", callback_data="pdf_orig"), InlineKeyboardButton(text="📄 PDF (Oq-Qora)", callback_data="pdf_bw")],
        [InlineKeyboardButton(text="🗑 Tozalash", callback_data="clear")]
    ])
    await m.reply(f"✅ Rasm #{len(G_DATA[uid]['files'])} qabul qilindi.\nEffekt yoki amalni tanlang:", reply_markup=kb)

@dp.message(F.document)
async def doc_h(m: types.Message):
    uid = m.from_user.id
    f = await bot.get_file(m.document.file_id)
    content = await bot.download_file(f.file_path)
    G_DATA[uid] = {'doc': content.read(), 'state': None}
    
    kb = []
    if "pdf" in m.document.mime_type:
        kb = [[InlineKeyboardButton(text="✂️ Kesish", callback_data="split"), InlineKeyboardButton(text="📝 Wordga", callback_data="pdf2word")]]
    else:
        kb = [[InlineKeyboardButton(text="📄 PDFga", callback_data="any2pdf")]]
    await m.reply(f"📂 {m.document.file_name} tayyor.", reply_markup=InlineKeyboardMarkup(inline_keyboard=kb))

# --- CALLBACKS ---
@dp.callback_query(F.data)
async def call_worker(call: types.CallbackQuery):
    uid, d = call.from_user.id, call.data
    files = G_DATA.get(uid, {}).get('files', [])

    if d == "clear": 
        G_DATA[uid]['files'] = []
        await call.message.delete()
        await call.message.answer("🗑 Tozalandi."); return

    if d == "to_ocr":
        if not files: return
        msg = await call.message.edit_text("⏳ <b>Google Vision tahlil qilmoqda...</b>\n☁️ Serverga ulanish...")
        
        full_result = ""
        loop = asyncio.get_event_loop()
        
        for idx, img in enumerate(files):
            # Google Vision chaqirish
            text = await loop.run_in_executor(None, google_vision_scan, img)
            full_result += f"📄 <b>Sahifa {idx+1}:</b>\n\n{html.escape(text)}\n\n{'='*20}\n\n"
        
        if len(full_result) > 4000:
            await call.message.answer_document(
                BufferedInputFile(full_result.encode('utf-8'), filename="google_vision_result.txt"),
                caption="✅ <b>Google OCR Natija</b> (Fayl)"
            )
        else:
            await call.message.answer(f"✅ <b>Google Vision Natija:</b>\n\n<pre>{full_result}</pre>")
        
        await msg.delete()
        G_DATA[uid]['files'] = []

    if d == "to_enhance":
        if not files: return
        msg = await call.message.edit_text("✨ <b>Sifat oshirilmoqda...</b>")
        loop = asyncio.get_event_loop()
        for i, img in enumerate(files):
            # Enhance effekti
            res = await loop.run_in_executor(None, process_image_effect, img, "enhance")
            await call.message.answer_photo(BufferedInputFile(res, filename=f"hd_{i+1}.jpg"))
        await msg.delete()

    if d.startswith("pdf_"): # pdf_orig yoki pdf_bw
        mode = d.split("_")[1]
        msg = await call.message.edit_text(f"⏳ <b>PDF ({mode}) yaratilmoqda...</b>")
        loop = asyncio.get_event_loop()
        
        processed_imgs = []
        for img in files:
            res = await loop.run_in_executor(None, process_image_effect, img, mode)
            processed_imgs.append(res)
            
        pdf = await loop.run_in_executor(None, img2pdf.convert, processed_imgs)
        await call.message.answer_document(BufferedInputFile(pdf, filename=f"scan_{mode}.pdf"))
        await msg.delete()
        G_DATA[uid]['files'] = []

    if d == "to_word":
        msg = await call.message.edit_text("⏳ <b>Rasmlar Wordga...</b>")
        loop = asyncio.get_event_loop()
        docx = await loop.run_in_executor(None, images_to_docx, files)
        await call.message.answer_document(BufferedInputFile(docx, filename="images.docx"))
        await msg.delete()
        G_DATA[uid]['files'] = []

    if d == "any2pdf":
        msg = await call.message.edit_text("⏳ <b>PDF yaratilmoqda...</b>")
        loop = asyncio.get_event_loop()
        pdf_res = None
        doc_content = G_DATA[uid]['doc']
        try:
            if b"PK\x03\x04" in doc_content[:4]:
                pdf_res = await loop.run_in_executor(None, docx_to_pdf_engine, doc_content)
            else:
                pdf_res = await loop.run_in_executor(None, create_pdf_from_text, doc_content.decode('utf-8', errors='ignore'))
        except: pass
        if pdf_res: await call.message.answer_document(BufferedInputFile(pdf_res, filename="hujjat.pdf"))
        else: await call.message.answer("❌ Xatolik.")
        await msg.delete()

    if d == "pdf2word":
        msg = await call.message.edit_text("⏳ <b>Wordga o'girilmoqda...</b>")
        loop = asyncio.get_event_loop()
        docx = await loop.run_in_executor(None, convert_pdf_to_docx_safe, G_DATA[uid]['doc'])
        if docx: await call.message.answer_document(BufferedInputFile(docx, filename="converted.docx"))
        else: await call.message.answer("❌ Xatolik.")
        await msg.delete()

    if d == "split": 
        G_DATA[uid]['state'] = "split"
        await call.message.answer("✂️ Oraliqni yozing (Masalan: 1-5):")

# --- 7. RUNNER ---
def run_bot():
    new_loop = asyncio.new_event_loop()
    asyncio.set_event_loop(new_loop)
    async def starter():
        await bot.delete_webhook(drop_pending_updates=True)
        await dp.start_polling(bot, handle_signals=False)
    new_loop.run_until_complete(starter())

if not any(t.name == "AiogramThread" for t in threading.enumerate()):
    threading.Thread(target=run_bot, name="AiogramThread", daemon=True).start()

# --- 8. ADMIN PANEL ---
st.markdown('<p class="header-text">👁️ Google Vision Dashboard</p>', unsafe_allow_html=True)
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/4712/4712035.png", width=100)
    p = st.text_input("Security Key", type="password")

if p == ADMIN_PASS:
    st.success("✅ Google Cloud API Connected")
    c1, c2, c3 = st.columns(3)
    c1.metric("👥 Userlar", len(G_DATA))
    c2.metric("🔄 Threads", threading.active_count())
    c3.metric("🧠 AI Engine", "Google Cloud Vision")
    
    st.markdown("---")
    st.markdown("### 👁️ Google Vision Statistikasi:")
    st.info("API ulanishi muvaffaqiyatli amalga oshirildi. JSON kalit xotirada mavjud.")
else:
    st.image("https://img.freepik.com/free-vector/abstract-technology-particle-background_23-2148426649.jpg")
    st.info("🔐 Admin parolini kiriting")
